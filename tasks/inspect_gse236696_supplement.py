"""Extract auditable text/table contents from the GSE236696 source-paper supplements.

This is a read-only inspection utility.  It intentionally preserves paragraph and
table order so that downstream biological claims can be traced to the original
supplement rather than to OCR or a secondary summary.
"""

from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from openpyxl import load_workbook


def extract_docx(path: Path) -> dict:
    document = Document(path)
    paragraphs = [
        {"index": index, "text": paragraph.text.strip()}
        for index, paragraph in enumerate(document.paragraphs)
        if paragraph.text.strip()
    ]
    tables = []
    for table_index, table in enumerate(document.tables):
        rows = []
        for row in table.rows:
            rows.append([cell.text.replace("\n", " / ").strip() for cell in row.cells])
        tables.append({"index": table_index, "rows": rows})
    return {
        "path": str(path),
        "paragraphs": paragraphs,
        "tables": tables,
    }


def extract_xlsx(path: Path) -> dict:
    workbook = load_workbook(path, data_only=False, read_only=True)
    sheets = []
    for worksheet in workbook.worksheets:
        rows = [
            [None if value is None else value for value in row]
            for row in worksheet.iter_rows(values_only=True)
        ]
        sheets.append(
            {
                "name": worksheet.title,
                "max_row": worksheet.max_row,
                "max_column": worksheet.max_column,
                "rows": rows,
            }
        )
    return {"path": str(path), "sheets": sheets}


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--supplement-dir",
        type=Path,
        default=Path("data/external/GSE236696/source_paper/supplementary"),
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("data/external/GSE236696/source_paper/supplement_contents.json"),
    )
    args = parser.parse_args()

    if not args.supplement_dir.is_dir():
        raise FileNotFoundError(args.supplement_dir)

    docx_files = sorted(args.supplement_dir.glob("*.docx"))
    xlsx_files = sorted(args.supplement_dir.glob("*.xlsx"))
    if not docx_files or not xlsx_files:
        raise RuntimeError("expected both DOCX and XLSX supplementary files")

    payload = {
        "status": "gse236696_supplement_inspection_complete",
        "docx": [extract_docx(path) for path in docx_files],
        "xlsx": [extract_xlsx(path) for path in xlsx_files],
    }
    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(json.dumps(payload, indent=2, ensure_ascii=False), encoding="utf-8")
    print(json.dumps({
        "status": payload["status"],
        "docx_files": len(payload["docx"]),
        "xlsx_files": len(payload["xlsx"]),
        "output": str(args.output),
    }, indent=2))


if __name__ == "__main__":
    main()
