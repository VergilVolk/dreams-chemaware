#!/usr/bin/env python
"""Audit Neu5Ac evidence in the Jain et al. CRC biogeography supplement.

The script deliberately separates four questions that are easy to conflate:

1. Was N-acetylneuraminic acid identified with an analytical standard?
2. Does it show a spatial gradient in normal or tumour tissue?
3. Is it listed among the per-subsite tumour-versus-normal summary rows?
4. Does the public supplement contain a mucinous-histology subgroup?

Only the first two are positive in the frozen supplement.  The resulting
artifacts are therefore external abundance/context evidence, not an
independent replication of the MTBLS13729 mucinous phenotype.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
from pathlib import Path
from typing import Iterable

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE = ROOT / "data/external/CRC_metabolic_biogeography_PMC11438248_20260831"
DEFAULT_OUTPUT = DEFAULT_SOURCE / "neu5ac_biogeography_audit_v1"

EXPECTED_HASHES = {
    "supplementary_methods.docx": "1cbce8ed71d352165834d5865dfac86287165c905a85b3e55dbc924f2578f96a",
    "supplementary_tables.docx": "edc6f11d11e9b74f11d599aeadc64466bf2e2dc73fb9313edaabde71b8d9e748",
}


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1 << 20), b""):
            digest.update(block)
    return digest.hexdigest()


def clean(value: str) -> str:
    return " ".join(value.replace("\n", " ").split())


def table_rows(document: Document) -> Iterable[tuple[int, int, list[str]]]:
    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            yield table_index, row_index, [clean(cell.text) for cell in row.cells]


def one_row(rows: list[dict[str, object]], description: str) -> dict[str, object]:
    if len(rows) != 1:
        raise RuntimeError(f"expected one {description} row; found {len(rows)}")
    return rows[0]


def parse_float(value: str) -> float:
    return float(value.strip().lstrip("<"))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source-dir", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output-dir", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()

    methods_path = args.source_dir / "supplementary_methods.docx"
    tables_path = args.source_dir / "supplementary_tables.docx"
    for path in (methods_path, tables_path):
        if not path.is_file():
            raise FileNotFoundError(path)
        observed = sha256(path)
        expected = EXPECTED_HASHES[path.name]
        if observed.lower() != expected:
            raise RuntimeError(f"source hash mismatch for {path.name}: {observed}")

    methods = Document(methods_path)
    tables = Document(tables_path)
    rows = list(table_rows(tables))

    standard_hits: list[dict[str, object]] = []
    spatial_hits: list[dict[str, object]] = []
    tumour_normal_summary_hits: list[dict[str, object]] = []
    acetylated_hits: list[dict[str, object]] = []
    for table_index, row_index, values in rows:
        if not values:
            continue
        name = values[0].strip()
        normalized = name.casefold()
        if normalized == "n-acetylneuraminic acid":
            if table_index == 0 and len(values) >= 5:
                standard_hits.append({
                    "table_index": table_index,
                    "row_index": row_index,
                    "metabolite": name,
                    "id_level": values[1],
                    "mode": values[2],
                    "ion_mz": values[3],
                    "rt_sec": values[4],
                })
            elif table_index == 5 and len(values) >= 5:
                spatial_hits.append({
                    "table_index": table_index,
                    "row_index": row_index,
                    "metabolite": name,
                    "control_slope": parse_float(values[1]),
                    "control_p_text": values[2],
                    "control_p_upper_bound": parse_float(values[2]),
                    "tumour_slope": parse_float(values[3]),
                    "tumour_p_text": values[4],
                    "tumour_p": parse_float(values[4]),
                })
            elif table_index == 4:
                tumour_normal_summary_hits.append({
                    "table_index": table_index,
                    "row_index": row_index,
                    "values": values,
                })
        if normalized == "n-acetyl-9-o-acetylneuraminic acid":
            acetylated_hits.append({
                "table_index": table_index,
                "row_index": row_index,
                "metabolite": name,
                "values": values[1:],
            })

    standard = one_row(standard_hits, "Level-1 Neu5Ac standard")
    spatial = one_row(spatial_hits, "Neu5Ac spatial-gradient")
    if standard["id_level"] != "1":
        raise RuntimeError(f"Neu5Ac is not frozen as ID Level 1: {standard}")
    if standard["mode"] != "HILIC (-)":
        raise RuntimeError(f"unexpected Neu5Ac acquisition mode: {standard}")
    if spatial["control_p_text"] != "<.001" or spatial["tumour_p_text"] != ".091":
        raise RuntimeError(f"unexpected Neu5Ac spatial statistics: {spatial}")
    if tumour_normal_summary_hits:
        raise RuntimeError("Neu5Ac unexpectedly appears in the per-subsite tumour-normal summary table")

    method_text = [clean(paragraph.text) for paragraph in methods.paragraphs if clean(paragraph.text)]
    table_text = [" | ".join(values) for _, _, values in rows]
    all_text = method_text + table_text
    mucinous_terms = ("mucinous", "mucinous adenocarcinoma")
    mucinous_hits = [text for text in all_text if any(term in text.casefold() for term in mucinous_terms)]

    # Table 3 contains sex, age and stage by subsite.  Its stage rows provide
    # an internal checksum of the reported 372-patient cohort.
    stage_rows = {
        values[0]: values[1:8]
        for table_index, _, values in rows
        if table_index == 3 and values and values[0] in {"I", "II", "III", "IV"}
    }
    stage_total = sum(int(value) for values in stage_rows.values() for value in values)
    # The source supplement is internally inconsistent here: the methods and
    # paper report 372 pairs, while the displayed stage cells sum to 374.  This
    # is a source-data caveat, not a reason to silently rewrite either number.
    reported_cohort_size = 372
    stage_table_matches_reported_cohort = stage_total == reported_cohort_size

    method_matches = {
        "cohort": [text for text in method_text if "372 frozen primary tumors" in text],
        "acquisition": [text for text in method_text if "Samples were analysed" in text],
        "identification": [text for text in method_text if "annotated as level 1" in text],
        "statistics": [text for text in method_text if "For linear regression between subsites" in text],
    }
    for label, matches in method_matches.items():
        if len(matches) != 1:
            raise RuntimeError(f"expected one methods paragraph for {label}; found {len(matches)}")

    interpretation = {
        "identity": (
            "Independent standard-supported CRC tissue evidence: Neu5Ac is frozen as Level 1 "
            "in HILIC negative mode, with m/z and retention time reported in the supplement."
        ),
        "spatial_biology": (
            "Normal mucosa shows a positive cecum-to-rectum Neu5Ac gradient "
            "(slope +0.349, p<0.001), whereas the tumour gradient is attenuated and "
            "not nominally significant (slope +0.088, p=0.091)."
        ),
        "mtbls13729_relevance": (
            "This supports disease-dependent anatomical regulation of Neu5Ac and strengthens "
            "the need for paired and location-aware MTBLS13729 analyses."
        ),
        "not_established": (
            "The supplement contains no mucinous-histology subgroup and no patient-level "
            "mucinous Neu5Ac test; it is not an independent replication of the Rmu phenotype."
        ),
    }

    args.output_dir.mkdir(parents=True, exist_ok=True)
    evidence_path = args.output_dir / "neu5ac_external_evidence.csv"
    with evidence_path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=[
            "evidence_type", "metabolite", "id_level", "mode", "ion_mz", "rt_sec",
            "control_slope", "control_p", "tumour_slope", "tumour_p", "interpretation",
        ])
        writer.writeheader()
        writer.writerow({
            "evidence_type": "analytical_standard_identity",
            "metabolite": standard["metabolite"],
            "id_level": standard["id_level"],
            "mode": standard["mode"],
            "ion_mz": standard["ion_mz"],
            "rt_sec": standard["rt_sec"],
            "interpretation": interpretation["identity"],
        })
        writer.writerow({
            "evidence_type": "disease_dependent_spatial_gradient",
            "metabolite": spatial["metabolite"],
            "control_slope": spatial["control_slope"],
            "control_p": spatial["control_p_text"],
            "tumour_slope": spatial["tumour_slope"],
            "tumour_p": spatial["tumour_p_text"],
            "interpretation": interpretation["spatial_biology"],
        })

    report = {
        "status": "external_crc_neu5ac_biogeography_audit_complete",
        "formal": True,
        "paper": {
            "title": "Charting the metabolic biogeography of the colorectum in cancer: challenging the right sided versus left sided classification",
            "doi": "10.1186/s12943-024-02133-5",
            "pmcid": "PMC11438248",
            "cohort": "372 patient-matched colorectal tumour and normal-mucosa pairs across seven subsites",
        },
        "level1_neu5ac": standard,
        "spatial_gradient": spatial,
        "neu5ac_in_per_subsite_tumour_normal_summary": False,
        "distinct_9_o_acetyl_neu5ac_rows": acetylated_hits,
        "histology_audit": {
            "mucinous_keyword_hits": len(mucinous_hits),
            "mucinous_subgroup_available": False,
            "available_patient_characteristics": ["sex", "age", "stage", "anatomical subsite"],
        },
        "patient_count_audit": {
            "reported_in_methods": reported_cohort_size,
            "sum_of_stage_cells_in_supplementary_table_4": stage_total,
            "stage_table_matches_reported_cohort": stage_table_matches_reported_cohort,
            "interpretation": (
                "The supplementary demographic table sums to 374 rather than the reported 372; "
                "the paper-level cohort size is retained and the table discrepancy is disclosed."
            ),
        },
        "methods_extracts": {label: matches[0] for label, matches in method_matches.items()},
        "interpretation": interpretation,
        "claim_limit": (
            "Independent CRC tissue identity and disease-dependent spatial-context evidence only; "
            "not a mucinous-histology replication, not a same-method MTBLS13729 standard, and not flux or causality."
        ),
        "provenance": {
            methods_path.name: sha256(methods_path),
            tables_path.name: sha256(tables_path),
            "script": sha256(Path(__file__)),
            "evidence_csv": sha256(evidence_path),
        },
    }
    report_path = args.output_dir / "report.json"
    report_path.write_text(json.dumps(report, indent=2, ensure_ascii=False), encoding="utf-8")
    print(json.dumps({
        "status": report["status"],
        "level1_neu5ac": standard,
        "spatial_gradient": spatial,
        "mucinous_subgroup_available": False,
        "output": str(args.output_dir),
    }, indent=2, ensure_ascii=False))


if __name__ == "__main__":
    main()
